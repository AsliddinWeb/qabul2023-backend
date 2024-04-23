from django.shortcuts import get_object_or_404
from drf_yasg.utils import swagger_auto_schema
from rest_framework import generics, permissions, status
from rest_framework.response import Response

from .models import Application
from .serializers import ApplicationSerializer
from shartnoma_app.models import Shartnoma

# Pdf
from django.conf import settings
import os
import docx
from datetime import datetime

from django.contrib.auth import get_user_model
User = get_user_model()

class ApplicationListCreateAPIView(generics.ListCreateAPIView):
    queryset = Application.objects.all()
    serializer_class = ApplicationSerializer
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        operation_summary="Application List",
        operation_description="Application List",
        tags=['Applications']
    )
    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_summary="Create Application",
        operation_description="Creates a new application and a corresponding contract.",
        tags=['Applications']
    )
    def post(self, request, *args, **kwargs):
        if not request.user.is_active:
            return Response({"detail": "Only active users can create passports."}, status=status.HTTP_403_FORBIDDEN)

        response = super().post(request, *args, **kwargs)
        if response.status_code == 201:
            application = Application.objects.get(id=response.data['id'])
            docx_path = self.generate_docx(application)
            Shartnoma.objects.create(
                user=application.user,
                apply=application,
                shartnoma_fayl=f"/shartnomalar/shartnoma_{application.id}.docx"
            )
        return response

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def generate_docx(self, application):
        # Ensure the directory exists
        docx_dir = os.path.join(settings.MEDIA_ROOT, "shartnomalar")
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        # Load and modify the DOCX template
        current_date = datetime.now().strftime('%d.%m.%Y')

        doc = docx.Document(os.path.join(settings.BASE_DIR, 'static/Shartnoma.docx'))
        for paragraph in doc.paragraphs:
            if '{{id}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{id}}', str(application.id))

            if '{{date}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{date}}', str(current_date))

            if '{{full_name1}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{full_name1}}', application.passport.get_full_name())

        for paragraph1 in doc.paragraphs:
            for paragraph in paragraph1.runs:
                if '{{talim_turi}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{talim_turi}}', str(application.talim_turi.nomi))
                    paragraph.bold = True

                if '{{talim_shakli}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{talim_shakli}}', str(application.talim_shakli.nomi))
                    paragraph.bold = True

                if '{{oqish_muddati}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{oqish_muddati}}',
                                                            str(application.yonalish.oqish_muddati))
                    paragraph.bold = True

                if '{{topshirgan_kursi}}' in paragraph.text:
                    if application.is_perevod:
                        paragraph.text = paragraph.text.replace('{{topshirgan_kursi}}',
                                                                application.perevod_diplom.topshiradigan_kursi)
                        paragraph.bold = True
                    else:
                        paragraph.text = paragraph.text.replace('{{topshirgan_kursi}}', "1-kurs")
                        paragraph.bold = True

                if '{{yonalish_kodi}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{yonalish_kodi}}',
                                                            str(application.yonalish.yonalish_kodi))
                    paragraph.bold = True

                if '{{yonalish}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{yonalish}}', str(application.yonalish.nomi))
                    paragraph.bold = True

                if '{{kontrakt_narxi}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{kontrakt_narxi}}',
                                                            str(application.yonalish.kontrakt_narxi))
                    paragraph.bold = True

                if '{{full_name}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{full_name}}', f"{application.passport.get_full_name()}")
                    paragraph.bold = True

                if '{{toliq_manzili}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{toliq_manzili}}',
                                                            f"{application.passport.get_full_address()}")
                    paragraph.bold = True

                if '{{passport_seriya}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{passport_seriya}}',
                                                            f"{application.passport.passport_number}")
                    paragraph.bold = True

                if '{{talaba_kodi}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{talaba_kodi}}', "Kod yo'q!")
                    paragraph.bold = True

                if '{{telefon_raqami}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{telefon_raqami}}', str(application.user.phone_number))
                    paragraph.bold = True

        # Save the modified DOCX file
        docx_path = os.path.join(docx_dir, f"shartnoma_{application.id}.docx")
        doc.save(docx_path)
        return docx_path

class ApplicationRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Application.objects.all()
    serializer_class = ApplicationSerializer
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        operation_summary="Retrieve Application",
        operation_description="Retrieves a specific application by ID.",
        tags=['Applications']
    )
    def get(self, request, *args, **kwargs):
        """ Retrieves a specific application. """
        return super().get(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_summary="Update Application",
        operation_description="Updates a specific application by ID. Supports full and partial updates.",
        tags=['Applications']
    )
    def put(self, request, *args, **kwargs):
        """ Updates a specific application. """
        return super().put(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_summary="Partially Update Application",
        operation_description="Partially updates specific fields of an application by ID.",
        tags=['Applications']
    )
    def patch(self, request, *args, **kwargs):
        """ Partially updates an application. """
        return super().patch(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_summary="Delete Application",
        operation_description="Deletes a specific application by ID.",
        tags=['Applications']
    )
    def delete(self, request, *args, **kwargs):
        """ Deletes a specific application. """
        return super().delete(request, *args, **kwargs)
